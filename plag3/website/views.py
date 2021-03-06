from django.shortcuts import render

# Create your views here.
from .models import user,assign,rep
from django.contrib import messages


# Create your views here.
def login(request):
    return render(request,'login.html')

def home(request):
    return render(request,'home.html')
def val(request):
    if request.method == 'POST':
        a = request.POST["roll"]
        b = request.POST["pass"]

        try:
            c = user.objects.get(rollno=a)
        except:
            print("in except user does not exist")
            messages.error(request,"roll no does not exist")
            return render(request, 'login.html')
        else:
            print("in success case")
            request.session['a'] = a
            if (c.password == b):
                messages.success(request, "Successfully loged into the account")
                return render(request, 'home.html')
            else:
                print("in rejection case wrong password")
                messages.error(request, "Invalid password")
                return render(request, 'login.html')

def uplo(request):
    if request.method == 'POST':
        fi = request.FILES['fil']
        print(request.session['a'])
        try:
            p = assign(roll=request.session['a'], file=fi)
            p.save()
        except:
            print("file has already been submitted")
            messages.error(request, "file was previously submitted")
            return render(request, 'login.html')
        else:
            messages.success(request, "Your file has been successfully uploaded")
            return render(request, 'login.html')

def report(request):
    calc()
    z = rep.objects.all()
    print(z)
    return render(request, 'report.html',{"z":z})

"""def calc():
    import os
    import docx
    import PyPDF2

    instance = rep.objects.all()
    instance.delete()

    def dicta(a):
        dict = {}
        for i in a:
            if dict.get(i) == None:
                dict[i] = 1
            else:
                dict[i] = dict[i] + 1
        return dict

    def sp(a):
        l = []
        lm = []
        for i in a:
            lm = i.split()
            for j in lm:
                # if (j.find(',') != -1) :
                j = j.replace("-", "")
                j = j.replace(' ', '')
                j = j.replace("\\n", " ")
                j = j.replace(',', '')
                j = j.replace('.', '')
                j = j.replace("'", "")
                j = j.lower()
                l.append(j)
        return l

    asd = []

    bp = "C:\\Users\\COMPACT\\Desktop\\plag3\\uploads\\upload\\am"
    # pp=int(input('Enter the plagarism % filter (0-99):'))
    s = ''
    for i in bp:
        if i == "\\":
            s = s + '/'
        else:
            s = s + i
    if s[-1] == '/':
        print(s)
    else:
        s = s + '/'
        print(s)

    basepath = s
    for entry in os.listdir(basepath):
        if os.path.isfile(os.path.join(basepath, entry)):
            asd.append(entry)
    # print(asd)
    pans = []
    pavg = []
    common = ['the', 'at', 'there', 'some', 'my', 'of', 'be', 'use', 'her', 'than', 'and', 'this', 'an', 'would',
              'first', 'a', 'have', 'each', 'make', 'water', 'to', 'from', 'which', 'like', 'been', 'in', 'or', 'she',
              'him', 'call', 'is', 'one', 'do', 'into', 'who', 'you', 'had', 'how', 'time', 'oil', 'that', 'by',
              'their', 'has', 'its', 'it', 'word', 'if', 'look', 'now', 'he', 'but', 'will', 'two', 'find', 'was',
              'not', 'up', 'more', 'long', 'for', 'what', 'other', 'write', 'down', 'on', 'all', 'about', 'go', 'day',
              'are', 'were', 'out', 'see', 'did', 'as', 'we', 'many', 'get', 'with', 'when', 'then', 'no', 'come',
              'his', 'your', 'them', 'way', 'made', 'they', 'can', 'these', 'could', 'may', 'I', 'said', 'so', 'part']
    if len(asd) >= 2:
        for i in range(len(asd)):
            doc1 = []
            doc1d = {}
            if (asd[i][-1:-5:-1][::-1] == 'docx'):
                d1 = docx.Document(basepath + asd[i])
                for p in d1.paragraphs:
                    doc1.append(p.text)
                doc1 = sp(doc1)
                doc1dog = dicta(doc1)
                doc1d = dicta(doc1)
            elif (asd[i][-1:-5:-1][::-1] == '.pdf'):
                d1 = open(basepath + asd[i], "rb")
                pdf_reader = PyPDF2.PdfFileReader(d1)
                # print(pdf_reader.numPages)   #to print no. of pages
                for p in range(pdf_reader.numPages):
                    pg1 = pdf_reader.getPage(p)  # first page starts from index zero
                    doc1.append(pg1.extractText())
                doc1 = sp(doc1)
                doc1dog = dicta(doc1)
                doc1d = dicta(doc1)

            else:
                continue
            plg = 0
            s = 0
            dog = 0
            lkghj = []
            for k in doc1d:
                s = s + doc1d[k]
            for j in range(len(asd)):
                if i != j:
                    # print(asd[i],'&',asd[j],'+++++++++++++')
                    doc2 = []
                    doc2d = {}
                    '''d2=docx.Document(basepath+asd[j])
                    for p in d2.paragraphs:
                        doc2.append(p.text)
                    doc2=sp(doc2)
                    doc2d=dicta(doc2)'''
                    if (asd[j][-1:-5:-1][::-1] == 'docx'):
                        d2 = docx.Document(basepath + asd[j])
                        for p in d2.paragraphs:
                            doc2.append(p.text)
                        doc2 = sp(doc2)
                        doc2d = dicta(doc2)
                    elif (asd[j][-1:-5:-1][::-1] == '.pdf'):
                        d2 = open(basepath + asd[j], "rb")
                        pdf_reader = PyPDF2.PdfFileReader(d2)
                        # print(pdf_reader.numPages)   #to print no. of pages
                        for p in range(pdf_reader.numPages):
                            pg1 = pdf_reader.getPage(p)  # first page starts from index zero
                            doc2.append(pg1.extractText())
                        doc2 = sp(doc2)
                        doc2d = dicta(doc2)

                    else:
                        continue
                    lkgh = []
                    if doc1 == doc2:
                        # print('********')
                        # print('Same',asd[i],'&',asd[j])
                        pans.append([asd[i], asd[j], 100, doc1])
                        break

                    else:

                        d = 0  # same words
                        for k in doc1dog:
                            if (k in doc2d) and (not (k in common)):
                                # print('y')
                                if (k not in lkgh):
                                    lkgh.append(k)
                                if doc2d[k] - doc1dog[k] <= 0:
                                    d = d + doc2d[k]

                                else:
                                    d = d + doc1dog[k]
                            # else:
                            # d=doc1d[k]
                        plg = d / s
                        pans.append([asd[i], asd[j], round(plg * 100, 2), lkgh])

                    for k in doc1d:
                        if (k in doc2d) and (not (k in common) and doc1d[k] > 0):
                            if (k not in lkghj):
                                lkghj.append(k)
                                # print('y')
                            if doc2d[k] - doc1d[k] <= 0:
                                doc1d[k] = doc1d[k] - doc2d[k]
                            else:
                                doc1d[k] = doc1d[k] - doc1d[k]
            dog = 0
            for wd in doc1dog:
                dog = dog + (doc1dog[wd] - doc1d[wd])
                # print(doc1dog,doc1d)
            pavg.append([asd[i], 'all', round((dog / s) * 100, 2), lkghj])

    print('Max % copied from files---------------------------')
    pmaxplg = []
    for i in asd:
        mxpg = 0
        mxj = 0
        for j in range(len(pans)):
            if i == pans[j][0]:
                if mxpg <= pans[j][2]:
                    mxpg = pans[j][2]
                    mxj = j
        pmaxplg.append(pans[mxj])
        # print('In ',pans[mxj][0],' from ',pans[mxj][1],' ---> ',pans[mxj][2],'%')
    # print('--------------------------------------------------')

    for i in pmaxplg:
        # print(i[0])
        for j in pavg:

            if i[0] == j[0]:
                if i[2] > j[2]:
                    print('In ', i[0], ' from ', i[1], ' ---> ', i[2],
                          '%-----------------------------------------------')
                    # print(i[3])
                    q = "upload/am/" + i[0]
                    p = assign.objects.get(file=q)
                    r = rep(rollno=p.roll, score=i[2])
                    print(p.roll,i[2])
                    r.save()

                else:
                    print('In ', j[0], ' from ', j[1], ' ---> ', j[2],
                          '%-----------------------------------------------')
                    # print(j[3])
                    q = "upload/am/" + j[0]
                    p = assign.objects.get(file=q)

                    r = rep(rollno=p.roll, score=j[2])
                    r.save()

                    # print(j[3])
    print('--------------------------------------------------') """


def calc():
    import os
    import docx
    import PyPDF2
    instance = rep.objects.all()
    instance.delete()

    def dicta(a):
        dict = {}
        for i in a:
            if dict.get(i) == None:
                dict[i] = 1
            else:
                dict[i] = dict[i] + 1
        return dict

    def sp(a):
        l = []
        lm = []
        for i in a:
            lm = i.split()
            for j in lm:
                # if (j.find(',') != -1) :
                j = j.replace("-", "")
                j = j.replace(' ', '')
                j = j.replace("\\n", " ")
                j = j.replace(',', '')
                j = j.replace('.', '')
                j = j.replace("'", "")
                j = j.lower()
                l.append(j)
        return l

    asd = []
    bp = "C:\\project\\plag3\\uploads\\upload\\am"
    # pp=int(input('Enter the plagarism % filter (0-99):'))
    s = ''
    for i in bp:
        if i == "\\":
            s = s + '/'
        else:
            s = s + i
    if s[-1] == '/':
        print(s)
    else:
        s = s + '/'
        print(s)

    basepath = s
    for entry in os.listdir(basepath):
        if os.path.isfile(os.path.join(basepath, entry)):
            asd.append(entry)
    # print(asd)
    pans = []
    pavg = []

    common = ['the', 'at', 'there', 'some', 'my', 'of', 'be', 'use', 'her', 'than', 'and', 'this', 'an', 'would',
              'first', 'a', 'have', 'each', 'make', 'water', 'to', 'from', 'which', 'like', 'been', 'in', 'or', 'she',
              'him', 'call', 'is', 'one', 'do', 'into', 'who', 'you', 'had', 'how', 'time', 'oil', 'that', 'by',
              'their', 'has', 'its', 'it', 'word', 'if', 'look', 'now', 'he', 'but', 'will', 'two', 'find', 'was',
              'not', 'up', 'more', 'long', 'for', 'what', 'other', 'write', 'down', 'on', 'all', 'about', 'go', 'day',
              'are', 'were', 'out', 'see', 'did', 'as', 'we', 'many', 'get', 'with', 'when', 'then', 'no', 'come',
              'his', 'your', 'them', 'way', 'made', 'they', 'can', 'these', 'could', 'may', 'I', 'said', 'so', 'part']
    if len(asd) >= 2:
        for i in range(len(asd)):
            doc1 = []
            doc1d = {}
            alli = []
            allf = []
            allp = []
            if (asd[i][-1:-5:-1][::-1] == 'docx'):
                d1 = docx.Document(basepath + asd[i])
                for p in d1.paragraphs:
                    doc1.append(p.text)
                doc1 = sp(doc1)
                doc1dog = dicta(doc1)
                doc1d = dicta(doc1)
            elif (asd[i][-1:-5:-1][::-1] == '.pdf'):
                d1 = open(basepath + asd[i], "rb")
                pdf_reader = PyPDF2.PdfFileReader(d1)
                # print(pdf_reader.numPages)   #to print no. of pages
                for p in range(pdf_reader.numPages):
                    pg1 = pdf_reader.getPage(p)  # first page starts from index zero
                    doc1.append(pg1.extractText())
                doc1 = sp(doc1)
                doc1dog = dicta(doc1)
                doc1d = dicta(doc1)

            else:
                continue
            plg = 0
            s = 0
            dog = 0
            lkghj = []

            for k in doc1d:
                s = s + doc1d[k]
            alld = s
            for j in range(len(asd)):
                if i != j:
                    # print(asd[i],'&',asd[j],'+++++++++++++')
                    doc2 = []
                    doc2d = {}
                    '''d2=docx.Document(basepath+asd[j])
                    for p in d2.paragraphs:
                        doc2.append(p.text)
                    doc2=sp(doc2)
                    doc2d=dicta(doc2)'''
                    if (asd[j][-1:-5:-1][::-1] == 'docx'):
                        d2 = docx.Document(basepath + asd[j])
                        for p in d2.paragraphs:
                            doc2.append(p.text)
                        doc2 = sp(doc2)
                        doc2d = dicta(doc2)
                    elif (asd[j][-1:-5:-1][::-1] == '.pdf'):
                        d2 = open(basepath + asd[j], "rb")
                        pdf_reader = PyPDF2.PdfFileReader(d2)
                        # print(pdf_reader.numPages)   #to print no. of pages
                        for p in range(pdf_reader.numPages):
                            pg1 = pdf_reader.getPage(p)  # first page starts from index zero
                            doc2.append(pg1.extractText())
                        doc2 = sp(doc2)
                        doc2d = dicta(doc2)

                    else:
                        continue
                    lkgh = []
                    if doc1 == doc2:
                        # print('********')
                        # print('Same',asd[i],'&',asd[j])
                        pans.append([asd[i], asd[j], 100, doc1])
                        break

                    else:

                        d = 0  # same words
                        for k in doc1dog:
                            if (k in doc2d) and (not (k in common)):
                                # print('y')
                                if (k not in lkgh):
                                    lkgh.append(k)
                                if doc2d[k] - doc1dog[k] <= 0:
                                    d = d + doc2d[k]

                                else:
                                    d = d + doc1dog[k]
                            # else:
                            # d=doc1d[k]
                        plg = d / s
                        pans.append([asd[i], asd[j], round(plg * 100, 2), lkgh])
                    fc = 0
                    for k in doc1d:
                        if (k in doc2d) and (not (k in common) and doc1d[k] > 0):
                            fc = 1
                            if (k not in lkghj):
                                lkghj.append(k)
                                # print('y')
                            if doc2d[k] - doc1d[k] <= 0:
                                doc1d[k] = doc1d[k] - doc2d[k]
                            else:
                                doc1d[k] = doc1d[k] - doc1d[k]

                    if fc == 1:
                        dog1 = 0
                        for wd in doc1d:
                            dog1 = dog1 + (doc1d[wd])

                        allf.append(asd[j])
                        allp.append(round(((alld - dog1) / s) * 100, 2))
                        alld = dog1
                        # print(asd[j])
                        fc = 0

            dog = 0
            for wd in doc1dog:
                dog = dog + (doc1dog[wd] - doc1d[wd])
                # print(doc1dog,doc1d)
            pavg.append([asd[i], allf, round((dog / s) * 100, 2), lkghj, allp])

    print('Max % copied from files---------------------------')
    pmaxplg = []
    for i in asd:
        mxpg = 0
        mxj = 0
        for j in range(len(pans)):
            if i == pans[j][0]:
                if mxpg <= pans[j][2]:
                    mxpg = pans[j][2]
                    mxj = j
        pmaxplg.append(pans[mxj])
        # print('In ',pans[mxj][0],' from ',pans[mxj][1],' ---> ',pans[mxj][2],'%')
    # print('--------------------------------------------------')

    for i in pmaxplg:
        # print(i[0])
        for j in pavg:

            if i[0] == j[0]:
                if i[2] > j[2]:
                    print('In ', i[0], ' from ', i[1], ' ---> ', i[2], '%------++++++++++++++')
                    struck ='<tr><td colspan=2>' + str(i[1]) + '-----' + str(i[2]) + ' % </td></tr>'
                    print(struck)
                    print(i[3])
                    q = "upload/am/" + i[0]
                    p = assign.objects.get(file=q)
                    r = rep(rollno=p.roll, score=i[2],mess=struck)
                    print(p.roll, i[2])
                    r.save()
                else:
                    print('In ', j[0], ' from all ---> ', j[2], '%------++++++++++++++')

                    struck = ""
                    for k in range(len(j[1])):
                        duk ='<tr><td colspan=2>' + str(j[1][k]) + "-----" + str(j[4][k]) + ' % </td></tr>'
                        struck = struck + str(duk)
                    print(struck)
                    q = "upload/am/" + j[0]
                    p = assign.objects.get(file=q)

                    r = rep(rollno=p.roll, score=j[2],mess=struck)
                    r.save()
                    print(j[3])

    print('--------------------------------------------------')







